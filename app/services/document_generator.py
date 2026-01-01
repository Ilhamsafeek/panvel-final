# File: app/services/document_generator.py
"""
Document Generator Service
Generates Word documents for correspondence
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import logging

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Service for generating Word documents"""
    
    @staticmethod
    def generate_correspondence_docx(
        content: str,
        subject: str = None,
        sender_name: str = None,
        recipient_name: str = None,
        reference: str = None
    ) -> BytesIO:
        """
        Generate a professional Word document for correspondence
        
        Args:
            content: Main content/body text
            subject: Subject line
            sender_name: Name of sender
            recipient_name: Name of recipient
            reference: Reference number
            
        Returns:
            BytesIO object containing the Word document
        """
        
        try:
            logger.info("📝 Generating Word document")
            
            # Create document
            doc = Document()
            
            # Set document margins (in inches)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add company header
            header_para = doc.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header_para.add_run('CALIM 360 - Smart Contract Lifecycle Management')
            header_run.bold = True
            header_run.font.size = Pt(14)
            header_run.font.color.rgb = RGBColor(39, 98, 203)  # #2762cb
            
            doc.add_paragraph()  # Empty line
            
            # Add date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f'Date: {datetime.now().strftime("%d %B %Y")}')
            date_run.font.size = Pt(11)
            
            # Add reference if provided
            if reference:
                ref_para = doc.add_paragraph()
                ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                ref_run = ref_para.add_run(f'Ref: {reference}')
                ref_run.font.size = Pt(11)
            
            doc.add_paragraph()  # Empty line
            
            # Add recipient if provided
            if recipient_name:
                recipient_para = doc.add_paragraph(f'To: {recipient_name}')
                recipient_para.paragraph_format.space_after = Pt(6)
            
            # Add subject if provided
            if subject:
                subject_para = doc.add_paragraph()
                subject_run = subject_para.add_run(f'Subject: {subject}')
                subject_run.bold = True
                subject_run.font.size = Pt(12)
                subject_para.paragraph_format.space_after = Pt(12)
            
            doc.add_paragraph()  # Empty line
            
            # Add main content
            # Split content by paragraphs and handle line breaks
            content = content.strip()
            
            # Split by double newlines first (paragraphs)
            paragraphs = content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Handle single newlines within paragraphs
                    lines = para_text.split('\n')
                    para = doc.add_paragraph()
                    
                    for i, line in enumerate(lines):
                        if line.strip():
                            if i > 0:
                                para.add_run('\n')  # Add line break for single newlines
                            run = para.add_run(line.strip())
                            run.font.size = Pt(11)
                            run.font.name = 'Calibri'
                    
                    para.paragraph_format.space_after = Pt(10)
                    para.paragraph_format.line_spacing = 1.15
            
            doc.add_paragraph()  # Empty line
            
            # Add sender signature if provided
            if sender_name:
                doc.add_paragraph()
                signature_para = doc.add_paragraph('Yours faithfully,')
                signature_para.paragraph_format.space_after = Pt(30)
                
                sender_para = doc.add_paragraph(sender_name)
                sender_run = sender_para.runs[0]
                sender_run.bold = True
            
            # Add footer
            doc.add_paragraph()
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(
                '_______________________________________________\n'
                'Generated by CALIM 360'
            )
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Save to BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            logger.info("✅ Word document generated successfully")
            return docx_buffer
            
        except Exception as e:
            logger.error(f"❌ Error generating Word document: {str(e)}")
            raise Exception(f"Failed to generate Word document: {str(e)}")
    
    @staticmethod
    def generate_simple_docx(content: str) -> BytesIO:
        """
        Generate a simple Word document with just content
        
        Args:
            content: Text content to include
            
        Returns:
            BytesIO object containing the Word document
        """
        
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add content
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para.paragraph_format.space_after = Pt(10)
                    for run in para.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Calibri'
            
            # Save to s
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            return docx_buffer
            
        except Exception as e:
            logger.error(f"❌ Error generating simple Word document: {str(e)}")
            raise Exception(f"Failed to generate Word document: {str(e)}")